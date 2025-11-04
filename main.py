from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse, FileResponse, Response, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import fitz  # PyMuPDF
import os
import re
import shutil
from docx import Document
import PyPDF2
from typing import List
import io
from langdetect import detect
import docx2txt

app = FastAPI(
    title="PDF Word Count and Redaction Service", 
    description="A service to manage PDF files, including word count, redaction, and classification."
)

# Middleware to allow CORS for testing on different environments
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "/tmp"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Define max file size (4MB to stay under Vercel's limit)
MAX_FILE_SIZE = 101 * 1024 * 1024  # 101MB in bytes

def count_words_in_pdf(pdf_path):
    """Counts words in a PDF, ignoring empty lines and comments."""
    try:
        doc = fitz.open(pdf_path)
        text = " ".join([page.get_text("text") for page in doc])
        # Remove comments (lines starting with # or //)
        text = re.sub(r"^\s*(#|//).*$", "", text, flags=re.MULTILINE)
        # Remove extra whitespace and count words
        words = re.findall(r"\b\w+\b", text)
        return len(words)
    except Exception as e:
        print(f"Error counting words in PDF: {e}")
        return None

def count_words_in_docx(docx_path):
    """Counts words in a DOCX file."""
    try:
        doc = Document(docx_path)
        text = " ".join([para.text for para in doc.paragraphs])
        words = re.findall(r"\b\w+\b", text)
        return len(words)
    except Exception as e:
        print(f"Error counting words in DOCX: {e}")
        return None

def count_words_in_text(text_path):
    """Counts words in a text file."""
    try:
        with open(text_path, "r") as file:
            text = file.read()
        words = re.findall(r"\b\w+\b", text)
        return len(words)
    except Exception as e:
        print(f"Error counting words in text file: {e}")
        return None

def validate_file_extension(file: UploadFile):
    """Validates the file extension to ensure it's one of the allowed types."""
    allowed_extensions = ['.pdf', '.docx', '.txt']
    file_extension = os.path.splitext(file.filename)[1].lower()
    if file_extension not in allowed_extensions:
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF, DOCX, and TXT files are allowed.")
    return file_extension

def count_words_in_pdf_bytes(pdf_bytes):
    """Counts words in a PDF from bytes, ignoring empty lines and comments."""
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = " ".join([page.get_text("text") for page in doc])
        # Remove comments (lines starting with # or //)
        text = re.sub(r"^\s*(#|//).*$", "", text, flags=re.MULTILINE)
        # Remove extra whitespace and count words
        words = re.findall(r"\b\w+\b", text)
        return len(words)
    except Exception as e:
        print(f"Error counting words in PDF: {e}")
        return None

def count_words_in_docx_bytes(docx_bytes):
    """Counts words in a DOCX file from bytes."""
    try:
        doc = Document(io.BytesIO(docx_bytes))
        text = " ".join([para.text for para in doc.paragraphs])
        words = re.findall(r"\b\w+\b", text)
        return len(words)
    except Exception as e:
        print(f"Error counting words in DOCX: {e}")
        return None

def count_words_in_text_bytes(text_bytes):
    """Counts words in text from bytes."""
    try:
        text = text_bytes.decode('utf-8')
        words = re.findall(r"\b\w+\b", text)
        return len(words)
    except Exception as e:
        print(f"Error counting words in text file: {e}")
        return None

@app.post("/count-words")
async def count_words_endpoint(file: UploadFile = File(...)):
    """Processes an uploaded file and counts the words in it."""
    try:
        print(f"Received file upload request: {file.filename}")
        
        # Check file size using Content-Length header if available
        content_length = None
        if hasattr(file, 'size') and file.size:
            content_length = file.size
            print(f"File size from attribute: {content_length} bytes")
        elif hasattr(file, 'headers'):
            content_length_header = file.headers.get('content-length')
            if content_length_header:
                try:
                    content_length = int(content_length_header)
                    print(f"File size from header: {content_length} bytes")
                except ValueError:
                    print("Could not parse content-length header")
                    pass

        # If size is still unknown, we'll check during streaming
        if content_length and content_length > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum allowed size is {MAX_FILE_SIZE/1024/1024}MB."
            )

        file_extension = validate_file_extension(file)
        print(f"File extension validated: {file_extension}")
        
        # Read file content into memory in chunks to monitor size
        file_content = b""
        chunk_size = 1024 * 1024  # 1MB chunks
        total_read = 0
        
        print("Starting to read file in chunks...")
        while True:
            try:
                chunk = await file.read(chunk_size)
                if not chunk:
                    break
                file_content += chunk
                total_read += len(chunk)
                print(f"Read chunk: {len(chunk)} bytes, Total: {total_read} bytes")
                
                # Check size during streaming
                if total_read > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=413,
                        detail=f"File too large. Maximum allowed size is {MAX_FILE_SIZE/1024/1024}MB."
                    )
            except Exception as e:
                print(f"Error reading chunk: {str(e)}")
                raise HTTPException(status_code=500, detail=f"Error reading file: {str(e)}")

        print(f"File reading complete. Total size: {total_read} bytes")
        
        # Process in memory instead of writing to disk
        word_count = None
        print(f"Processing file as {file_extension}...")
        if file_extension == '.pdf':
            word_count = count_words_in_pdf_bytes(file_content)
        elif file_extension == '.docx':
            word_count = count_words_in_docx_bytes(file_content)
        elif file_extension == '.txt':
            word_count = count_words_in_text_bytes(file_content)

        if word_count is None:
            raise HTTPException(status_code=500, detail="Error processing the file.")

        print(f"Word count: {word_count}")
        response = {
            "total_words": word_count
        }

        return JSONResponse(content=response)
    except HTTPException:
        raise
    except Exception as e:
        print(f"Unexpected error in count_words_endpoint: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")

def redact_submission_ids_bytes(pdf_bytes):
    """Redacts Submission IDs and 'Document Details' from a PDF in memory."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    for page_num, page in enumerate(doc):
        # Search for and redact Submission ID
        text_instances = page.search_for("Submission ID trn:oid:::")
        for inst in text_instances:
            rect = fitz.Rect(inst.x0, inst.y0, inst.x1 + 100, inst.y1)
            # Add redaction annotation instead of drawing white rectangle
            page.add_redact_annot(rect, fill=(1, 1, 1))
        
        if page_num == 0:
            # Search for and redact Document Details
            details_instances = page.search_for("Document Details")
            for inst in details_instances:
                rect = fitz.Rect(0, inst.y0 - 50, page.rect.x1, inst.y0)
                # Add redaction annotation instead of drawing white rectangle
                page.add_redact_annot(rect, fill=(1, 1, 1))
        
        # Apply redactions on each page
        page.apply_redactions()

    # Save to bytes instead of file
    output_bytes = doc.tobytes()
    doc.close()
    return output_bytes

@app.post("/redact")
async def redact_pdf(file: UploadFile = File(...)):
    """Redacts sensitive information from an uploaded PDF and returns the file directly."""
    # Read file content in chunks
    file_content = b""
    chunk_size = 1024 * 1024  # 1MB chunks
    total_read = 0
    
    while True:
        chunk = await file.read(chunk_size)
        if not chunk:
            break
        file_content += chunk
        total_read += len(chunk)
        
        if total_read > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum allowed size is {MAX_FILE_SIZE/1024/1024}MB."
            )

    # Process in memory
    redacted_bytes = redact_submission_ids_bytes(file_content)
    
    return StreamingResponse(
        io.BytesIO(redacted_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=redacted_{file.filename}"}
    )

def extract_second_page_text_bytes(pdf_bytes):
    """Extracts text from the second page of a PDF from bytes."""
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if len(doc) < 2:
            doc.close()
            return None
        text = doc[1].get_text("text")
        doc.close()
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None

def classify_pdf_bytes(pdf_bytes):
    """Classifies a PDF based on plagiarism and AI detection patterns from bytes."""
    text = extract_second_page_text_bytes(pdf_bytes)
    if not text:
        return {"error": "PDF does not have a second page or could not be read."}

    # Check for special case: "10… Overall Similarity" should be treated as 100%
    special_similarity_match = re.search(r"10…\s*Overall Similarity", text)
    similarity_match = re.search(r"(\d+)%\s*Overall Similarity", text)
    ai_match = re.search(r"(?:(\d+)\*?%|(\*%))(?:\s*detected as AI)", text)

    result = {
        "type": "Unknown",
        "Overall Similarity": None,
        "AI Detection": None,
        "AI Detection Asterisk": False,
        "Below_Threshold": False,
        "special_character_found": False
    }

    # Determine the similarity value
    similarity_value = None
    if special_similarity_match:
        similarity_value = 100  # Special case: "10… Overall Similarity" = 100%
        result["special_character_found"] = True
    elif similarity_match:
        similarity_value = int(similarity_match.group(1))

    if similarity_value is not None and ai_match:
        result["type"] = "Plagiarism and AI Detection Report"
        result["Overall Similarity"] = similarity_value
        if ai_match.group(1):
            result["AI Detection"] = int(ai_match.group(1))
            result["AI Detection Asterisk"] = '*' in ai_match.group(0)
        else:
            result["AI Detection"] = -1  # "<20" AI detection case
            result["Below_Threshold"] = True
            result["AI Detection Asterisk"] = True
    elif similarity_value is not None:
        result["type"] = "Plagiarism Report"
        result["Overall Similarity"] = similarity_value
    elif ai_match:
        result["type"] = "AI Detection Report"
        if ai_match.group(1):
            result["AI Detection"] = int(ai_match.group(1))
            result["AI Detection Asterisk"] = '*' in ai_match.group(0)
        else:
            result["AI Detection"] = -1
            result["Below_Threshold"] = True
            result["AI Detection Asterisk"] = True

    return result

def fix_similarity_percentage_bytes(pdf_bytes):
    """Replaces '10… Overall Similarity' with '100% Overall Similarity' in a PDF from bytes."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    
    for page_num, page in enumerate(doc):
        # Search for the text "10… Overall Similarity"
        text_instances = page.search_for("10… Overall Similarity")
        
        for inst in text_instances:
            # Get the text rectangle
            rect = fitz.Rect(inst.x0, inst.y0, inst.x1, inst.y1)
            
            # Draw a white rectangle to cover the old text
            page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
            
            # Insert the new text "100% Overall Similarity"
            # Position the text at the same location
            text_point = fitz.Point(inst.x0, inst.y1 - 2)  # Slight adjustment for text positioning
            page.insert_text(text_point, "100% Overall Similarity", fontsize=10, color=(0, 0, 0))
    
    output_bytes = doc.tobytes()
    doc.close()
    return output_bytes

@app.post("/fix-similarity")
async def fix_similarity_endpoint(file: UploadFile = File(...)):
    """Fixes '10… Overall Similarity' to '100% Overall Similarity' in a PDF and returns the modified file."""
    # Validate file extension
    file_extension = validate_file_extension(file)
    if file_extension != '.pdf':
        raise HTTPException(status_code=400, detail="Only PDF files are supported for similarity fixing.")
    
    # Read file content in chunks
    file_content = b""
    chunk_size = 1024 * 1024  # 1MB chunks
    total_read = 0
    
    while True:
        chunk = await file.read(chunk_size)
        if not chunk:
            break
        file_content += chunk
        total_read += len(chunk)
        
        if total_read > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum allowed size is {MAX_FILE_SIZE/1024/1024}MB."
            )
    
    # Fix the similarity percentage in memory
    fixed_bytes = fix_similarity_percentage_bytes(file_content)
    
    return StreamingResponse(
        io.BytesIO(fixed_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=fixed_{file.filename}"}
    )

@app.post("/classify")
async def classify_pdf_endpoint(file: UploadFile = File(...)):
    """Processes an uploaded PDF and classifies it based on AI detection and plagiarism patterns."""
    # Read file content in chunks
    file_content = b""
    chunk_size = 1024 * 1024  # 1MB chunks
    total_read = 0
    
    while True:
        chunk = await file.read(chunk_size)
        if not chunk:
            break
        file_content += chunk
        total_read += len(chunk)
        
        if total_read > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum allowed size is {MAX_FILE_SIZE/1024/1024}MB."
            )

    result = classify_pdf_bytes(file_content)
    return JSONResponse(content=result)

@app.get("/")
async def read_root():
    """Root endpoint that confirms the service is running and provides route information."""
    return JSONResponse({
        "message": "PDF and Document Service is online. Access /docs for the API documentation.",
        "routes": [
            {"path": "/count-words", "description": "Count words in a document"},
            {"path": "/redact", "description": "Redact sensitive information from a PDF"},
            {"path": "/classify", "description": "Classify a PDF for AI and plagiarism detection"},
            {"path": "/fix-similarity", "description": "Fix '10… Overall Similarity' to '100% Overall Similarity' in a PDF"},
            {"path": "/status", "description": "Get status of processed files"},
            {"path": "/merge-files", "description": "Merge multiple PDF or DOCX files into a single file"},
            {"path": "/check-english", "description": "Check the percentage of English content in a file"},
            {"path": "/check-language", "description": "Check the percentage of supported language content (English, Spanish, Japanese) in a file"},
            {"path": "/test-upload", "description": "Test endpoint for file upload debugging"}
        ]
    })

@app.post("/test-upload")
async def test_upload_endpoint(file: UploadFile = File(...)):
    """Test endpoint to diagnose file upload issues."""
    try:
        result = {
            "filename": file.filename,
            "content_type": file.content_type,
            "headers": dict(file.headers) if hasattr(file, 'headers') else {},
            "size_attribute": getattr(file, 'size', None),
            "chunks_received": 0,
            "total_bytes": 0,
            "status": "success"
        }
        
        # Try to read the file in chunks
        chunk_size = 1024 * 1024  # 1MB
        chunk_count = 0
        
        while True:
            try:
                chunk = await file.read(chunk_size)
                if not chunk:
                    break
                chunk_count += 1
                result["total_bytes"] += len(chunk)
                result["chunks_received"] = chunk_count
                print(f"Test upload: Received chunk {chunk_count}, size: {len(chunk)} bytes, total: {result['total_bytes']} bytes")
            except Exception as e:
                result["status"] = f"error_at_chunk_{chunk_count}"
                result["error"] = str(e)
                print(f"Test upload error at chunk {chunk_count}: {str(e)}")
                break
        
        print(f"Test upload complete: {result}")
        return JSONResponse(content=result)
    except Exception as e:
        print(f"Test upload endpoint error: {str(e)}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={"status": "error", "error": str(e), "traceback": traceback.format_exc()}
        )

@app.get("/status")
async def get_status():
    """Returns a list of files and their status in the upload directory."""
    try:
        files = os.listdir(UPLOAD_DIR)
        if not files:
            return JSONResponse(status_code=200, content={"message": "No files found in the upload directory."})

        file_list = [{"filename": file, "status": "redacted" if file.startswith("redacted_") else "uploaded"} for file in files]
        return JSONResponse(status_code=200, content={"files": file_list})

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

def merge_pdfs_in_memory(pdf_files: List[UploadFile]) -> bytes:
    """Merge multiple PDF files in memory."""
    merger = PyPDF2.PdfMerger()
    try:
        for pdf_file in pdf_files:
            # Read the file content
            content = pdf_file.file.read()
            # Create a BytesIO object with the content
            pdf_stream = io.BytesIO(content)
            merger.append(pdf_stream)
            # Reset the file pointer for the next iteration
            pdf_file.file.seek(0)
        
        output = io.BytesIO()
        merger.write(output)
        merger.close()
        return output.getvalue()
    except Exception as e:
        print(f"Error merging PDFs: {str(e)}")
        return None
    finally:
        merger.close()

def merge_docx_in_memory(docx_files: List[UploadFile]) -> bytes:
    """Merge multiple DOCX files in memory."""
    try:
        merged_doc = Document()
        
        for index, file in enumerate(docx_files):
            # Read the file content
            content = file.file.read()
            # Create a BytesIO object with the content
            doc_stream = io.BytesIO(content)
            doc = Document(doc_stream)
            
            # Copy all elements from the source document
            for element in doc.element.body:
                merged_doc.element.body.append(element)
            
            # Add page break between documents (except after the last one)
            if index != len(docx_files) - 1:
                merged_doc.add_page_break()
            
            # Reset the file pointer for the next iteration
            file.file.seek(0)
        
        # Save to BytesIO
        output = io.BytesIO()
        merged_doc.save(output)
        return output.getvalue()
    except Exception as e:
        print(f"Error merging DOCX files: {str(e)}")
        return None

@app.post("/merge-files")
async def merge_files(files: List[UploadFile] = File(...)):
    """Merge multiple PDF or DOCX files into a single file."""
    try:
        if len(files) < 2:
            raise HTTPException(status_code=400, detail="At least 2 files are required for merging.")
        
        # Check if all files are of the same type
        file_extensions = [os.path.splitext(file.filename)[1].lower() for file in files]
        if not all(ext == file_extensions[0] for ext in file_extensions):
            raise HTTPException(status_code=400, detail="All files must be of the same type (PDF or DOCX).")
        
        file_type = file_extensions[0]
        if file_type not in ['.pdf', '.docx']:
            raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported for merging.")
        
        # Merge files based on type
        merged_content = None
        if file_type == '.pdf':
            merged_content = merge_pdfs_in_memory(files)
        else:  # .docx
            merged_content = merge_docx_in_memory(files)
        
        if not merged_content:
            raise HTTPException(status_code=500, detail="Error merging files. Please check the file formats and try again.")
        
        # Create a BytesIO object with the merged content
        file_stream = io.BytesIO(merged_content)
        
        return StreamingResponse(
            file_stream,
            media_type="application/pdf" if file_type == '.pdf' else "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=merged_{files[0].filename}"
            }
        )
    except Exception as e:
        print(f"Error in merge_files endpoint: {str(e)}")
        raise HTTPException(status_code=500, detail=f"An error occurred while processing the files: {str(e)}")

def extract_pdf_content(pdf_file):
    """Extract text content from a PDF file."""
    try:
        print("Starting PDF content extraction...")
        reader = PyPDF2.PdfReader(pdf_file)
        print(f"Number of pages in PDF: {len(reader.pages)}")
        
        content = []
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
                print(f"Page {i+1} text length: {len(text) if text else 0}")
                if text and text.strip():
                    content.append(text.strip())
            except Exception as page_error:
                print(f"Error extracting text from page {i+1}: {str(page_error)}")
                continue
        
        if not content:
            print("No text content extracted from PDF")
            return ""
            
        full_content = ' '.join(content)
        print(f"Total extracted text length: {len(full_content)}")
        return full_content
    except Exception as e:
        print(f"Error in extract_pdf_content: {str(e)}")
        return ""

def extract_docx_content(docx_file):
    """Extract text content from a DOCX file."""
    try:
        doc = Document(docx_file)
        content = []
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        return ' '.join(content) if content else ""
    except Exception as e:
        print(f"Error in extract_docx_content: {str(e)}")
        return ""

def extract_doc_content(doc_file):
    """Extract text content from a DOC file."""
    try:
        content = docx2txt.process(doc_file)
        return content if content else ""
    except Exception as e:
        print(f"Error in extract_doc_content: {str(e)}")
        return ""

def check_language(content, target_languages=['en', 'es', 'ja']):
    """Check if content is in one of the supported languages (English, Spanish, Japanese)."""
    try:
        if not content or len(content.strip()) < 10:  # Minimum content length check
            return False, 'unknown'
        lang = detect(content)
        return lang in target_languages, lang
    except Exception as e:
        print(f"Error in check_language: {str(e)}")
        return False, 'unknown'

async def process_file(file: UploadFile):
    """Process a file and return its content parts."""
    try:
        print(f"Processing file: {file.filename}")
        file_extension = os.path.splitext(file.filename)[1].lower()
        print(f"File extension: {file_extension}")
        content = ""

        # Read file content into memory in chunks
        file_content = b""
        chunk_size = 1024 * 1024  # 1MB chunks
        total_read = 0
        
        while True:
            chunk = await file.read(chunk_size)
            if not chunk:
                break
            file_content += chunk
            total_read += len(chunk)
            
            if total_read > MAX_FILE_SIZE:
                raise ValueError(f"File too large. Maximum allowed size is {MAX_FILE_SIZE/1024/1024}MB.")
        
        print(f"File size: {len(file_content)} bytes")
        if not file_content:
            raise ValueError("Empty file content")
        
        # Create BytesIO object
        file_stream = io.BytesIO(file_content)
        
        if file_extension == ".pdf":
            content = extract_pdf_content(file_stream)
        elif file_extension == ".docx":
            content = extract_docx_content(file_stream)
        elif file_extension == ".doc":
            content = extract_doc_content(file_stream)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

        print(f"Extracted content length: {len(content) if content else 0}")
        if not content or len(content.strip()) < 50:  # Minimum content length check
            raise ValueError("No sufficient content could be extracted from the file")

        # Split content into parts, ensuring minimum part size
        content_length = len(content)
        print(f"Content length before splitting: {content_length}")
        
        if content_length < 300:  # If content is too short, just use it as one part
            content_parts = [content]
        else:
            part_size = max(content_length // 30, 100)  # Minimum 100 characters per part
            content_parts = [content[i:i+part_size] for i in range(0, content_length, part_size)]
        
        # Filter out empty or too small parts
        content_parts = [part for part in content_parts if len(part.strip()) >= 10]
        print(f"Number of content parts after filtering: {len(content_parts)}")
        
        if not content_parts:
            raise ValueError("No valid content parts could be created")
            
        return content_parts
    except Exception as e:
        print(f"Error in process_file: {str(e)}")
        raise

@app.post("/check-language")
async def check_language_percentage(file: UploadFile = File(...)):
    """Check the percentage of supported language content (English, Spanish, Japanese) in a file."""
    try:
        print(f"Starting language check for file: {file.filename}")
        content_parts = await process_file(file)
        print(f"Number of content parts to analyze: {len(content_parts)}")
        
        if not content_parts:
            raise ValueError("No content parts to analyze")
            
        supported_language_count = 0
        total_parts = len(content_parts)
        language_breakdown = {'en': 0, 'es': 0, 'ja': 0, 'other': 0}
        
        for i, part in enumerate(content_parts):
            try:
                print(f"Analyzing part {i+1}/{total_parts} (length: {len(part)})")
                is_supported, detected_lang = check_language(part)
                if is_supported:
                    supported_language_count += 1
                    if detected_lang in language_breakdown:
                        language_breakdown[detected_lang] += 1
                    else:
                        language_breakdown['other'] += 1
                else:
                    language_breakdown['other'] += 1
                print(f"Part {i+1}/{total_parts}: {'Supported language (' + detected_lang + ')' if is_supported else 'Not supported language (' + detected_lang + ')'}")
            except Exception as part_error:
                print(f"Error processing part {i+1}: {str(part_error)}")
                language_breakdown['other'] += 1
                continue

        # Calculate percentage of supported language parts
        supported_percentage = (supported_language_count / total_parts) * 100 if total_parts > 0 else 0
        print(f"Final results - Supported language parts: {supported_language_count}/{total_parts} ({supported_percentage:.2f}%)")
        
        return JSONResponse(
            status_code=200,
            content={
                "filename": file.filename,
                "supported_language_percentage": round(supported_percentage, 2),
                "total_parts_analyzed": total_parts,
                "supported_language_parts": supported_language_count,
                "language_breakdown": {
                    "english": language_breakdown['en'],
                    "spanish": language_breakdown['es'],
                    "japanese": language_breakdown['ja'],
                    "other": language_breakdown['other']
                }
            }
        )
    except Exception as e:
        print(f"Error in check_language_percentage: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

@app.post("/check-english")
async def check_english_percentage(file: UploadFile = File(...)):
    """Check the percentage of English content in a file."""
    try:
        print(f"Starting English check for file: {file.filename}")
        content_parts = await process_file(file)
        print(f"Number of content parts to analyze: {len(content_parts)}")
        
        if not content_parts:
            raise ValueError("No content parts to analyze")
            
        english_count = 0
        total_parts = len(content_parts)
        
        for i, part in enumerate(content_parts):
            try:
                print(f"Analyzing part {i+1}/{total_parts} (length: {len(part)})")
                is_supported, detected_lang = check_language(part, target_languages=['en'])
                if is_supported:
                    english_count += 1
                print(f"Part {i+1}/{total_parts}: {'English' if is_supported else 'Not English'}")
            except Exception as part_error:
                print(f"Error processing part {i+1}: {str(part_error)}")
                continue

        # Calculate percentage of English parts
        english_percentage = (english_count / total_parts) * 100 if total_parts > 0 else 0
        print(f"Final results - English parts: {english_count}/{total_parts} ({english_percentage:.2f}%)")
        
        return JSONResponse(
            status_code=200,
            content={
                "filename": file.filename,
                "english_percentage": round(english_percentage, 2),
                "total_parts_analyzed": total_parts,
                "english_parts": english_count
            }
        )
    except Exception as e:
        print(f"Error in check_english_percentage: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")